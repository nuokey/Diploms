import docx
from docx.shared import Inches
from random import choice
import mimesis

person = mimesis.Person('ru')

uslugge = ['лучшую прическу', "лучшую улыбку", "худшие отметки", "лучшие отметки", "лучший диплом"]
images = ['man.jpeg', 'woman.jpg']

doc = docx.Document()

def make_diplom():
    global doc
    doc.add_heading('Диплом', 0)

    image = choice(images)
    doc.add_picture(image, width=Inches(6))

    doc.add_heading(f'Награждается за {choice(uslugge)}', 0)
    doc.add_paragraph(person.first_name() + ' ' + person.last_name())

    doc.add_paragraph('Директор: Ганицев Тимофей')

    if image == 'man.jpeg':
        range_i = 9
    else:
        range_i = 4

    for i in range(range_i):
        doc.add_paragraph(' ')

for i in range(50):
    make_diplom()

doc.save('new_doc.docx')